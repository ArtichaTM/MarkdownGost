from docx.enum.text import WD_COLOR_INDEX

from . import logger
from ._flags import MacrosFlags
from ._mixins import DuringDocxCreation


class Macros(DuringDocxCreation):
    """Changes background color of run. Use WD_COLOR_INDEX names"""
    __slots__ = ()

    def __init__(self, macros) -> None:
        super().__init__(macros)
        self.macros = macros

    def process_during_docx_creation(self, p, context):
        if len(self.macros.args) != 1:
            logger.info(
                f'Макрос "{self.get_name()}" требует хотя бы 1 аргумент'
            )
            return [p.add_run("<Ошибка аргументов>")]
        name = self.macros.args[0]
        names = {i.name for i in WD_COLOR_INDEX}
        if name not in names:
            logger.info(
                f'Макрос "{self.get_name()}": второй аргумент '
                f' требует значение из WD_COLOR_INDEX, но '
                f'{name[:30]} не элемент {names}'
            )
            return [p.add_run("<Ошибка аргументов>")]
        run = p.add_run(self.macros.value)
        run.font.highlight_color = getattr(WD_COLOR_INDEX, self.macros.args[0])

        return [run]

    @staticmethod
    def flags():
        return MacrosFlags.NONE
